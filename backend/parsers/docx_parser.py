"""DOCX text extraction using python-docx."""

from __future__ import annotations

from pathlib import Path

from backend.models.ingestion import IngestedDocument, TextBlock
from backend.parsers.text_utils import clean_text_block


def extract_docx(file_path: Path) -> IngestedDocument:
    """Extract structured text from a DOCX file."""
    from docx import Document

    doc = Document(str(file_path))
    blocks: list[TextBlock] = []
    full_text_parts: list[str] = []
    warnings: list[str] = []
    block_index = 0
    char_offset = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        cleaned = clean_text_block(text)
        if not cleaned:
            continue

        # Detect headings from paragraph style
        is_heading = False
        heading_level = None
        style_name = (para.style.name or "").lower()
        if "heading" in style_name:
            is_heading = True
            # Try to extract level from style name like "Heading 1", "Heading 2"
            for ch in style_name:
                if ch.isdigit():
                    heading_level = int(ch)
                    break
            if heading_level is None:
                heading_level = 3

        tb = TextBlock(
            text=cleaned,
            page_number=None,
            block_index=block_index,
            is_heading=is_heading,
            heading_level=heading_level,
            char_offset_start=char_offset,
            char_offset_end=char_offset + len(cleaned),
        )
        blocks.append(tb)
        full_text_parts.append(cleaned)
        char_offset += len(cleaned) + 1
        block_index += 1

    full_text = "\n".join(full_text_parts)

    return IngestedDocument(
        filename=file_path.name,
        file_type="docx",
        total_pages=None,
        blocks=blocks,
        full_text=full_text,
        warnings=warnings,
    )
